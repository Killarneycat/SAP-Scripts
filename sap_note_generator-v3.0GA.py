"""Script to help you conveniently create SAP notes"""
"""creation by Michael Karek"""
"""support by Anne Kirchhof"""
"""An idea by Lorenzo de Luca and Michael Karek"""
"""please contact Michael Karek for questions"""
"""via telegram @diablescat"""
"""This is the GA Release 3.0 GA"""

import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from tkinter.scrolledtext import ScrolledText
from tkinter.filedialog import askopenfilename, asksaveasfilename
import re
from datetime import datetime
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas

def _on_mousewheel(event):
    canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

def clearTextInput():
    entry_case_number.delete(0, tk.END)
    text_issue_description_cx.delete("1.0", "end")
    text_issue_description_tse.delete("1.0", "end")
    text_entry_error.delete("1.0", "end")
    entry_customer_name.delete(0, tk.END)
    entry_customer_phone.delete(0, tk.END)
    entry_customer_language.delete(0, tk.END)
    entry_product_typ.delete(0, tk.END)
    entry_product_sn.delete(0, tk.END)
    entry_firmware_version.delete(0, tk.END)
    entry_ha_count.delete(0, tk.END)
    entry_access_id_fw.delete(0, tk.END)
    entry_access_id_expired.delete(0, tk.END)
    entry_access_id_fw2.delete(0, tk.END)
    entry_access_id_expired3.delete(0, tk.END)
    entry_access_id_uid.delete(0, tk.END)
    entry_access_id_expired2.delete(0, tk.END)
    entry_sophos_case_ftp_username.delete(0, tk.END)
    entry_sophos_case_ftp_password.delete(0, tk.END)
    entry_frequency_of_the_issue.delete(0, tk.END)
    entry_text_repro_possible.delete("1.0", "end")
    entry_timestamp.delete(0, tk.END)
    text_source_details.delete("1.0", "end")
    entry_workaround.delete("1.0", "end")
    combo_costumer_status.set("")
    combo_new_configuration.set("")
    combo_product_cat.set("")
    combo_ha.set("")
    combo_mode_ha.set("")
    combo_repro.set("")
    combo_workaround.set("")
    combo_poa.set("")
    combo_log_ftp.set("")
    combo_log_ss.set("")
    combo_log_a.set("")
    checkbox1_var.set(False)
    checkbox2_var.set(False)
    checkbox3_var.set(False)
    checkbox4_var.set(False)

def validate_int_input(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    if text in "0123456789":
        return True
    elif text == "":
        return True
    else:
        return False

def validate_version_input(action, index, value_if_allowed, prior_value, text, validation_type, trigger_type, widget_name):
    if re.match(r'^[0-9.]*$', value_if_allowed):
        return True
    else:
        return False

def submit():
    if not entry_case_number.get().isdigit() or int(entry_case_number.get()) <= 0:
        messagebox.showerror("Error", "Case Number is required.\nCase Number must be a positive integer.")
        return

    case_number = entry_case_number.get()
    if not case_number:
        messagebox.showerror("Error", "Case Number is required.")
        return

    issue_description_cx = text_issue_description_cx.get("1.0", tk.END).strip()
    issue_description_tse = text_issue_description_tse.get("1.0", tk.END).strip()
    error = text_entry_error.get("1.0", tk.END).strip()
    customer_status = combo_costumer_status.get()
    customer_name = entry_customer_name.get().strip()
    customer_phone = entry_customer_phone.get().strip()
    customer_language = entry_customer_language.get().strip()
    new_configuration = combo_new_configuration.get()
    product_cat = combo_product_cat.get()
    product_formatted = product_cat.lower()
    product_typ = entry_product_typ.get().strip()
    product_formatted_type = product_typ.lower()
    product_sn = entry_product_sn.get().strip()
    firmware_version = entry_firmware_version.get().strip()
    ha = combo_ha.get()
    ha_mode = combo_mode_ha.get()
    ha_count = entry_ha_count.get().strip()
    access_id_fw = entry_access_id_fw.get().strip()
    access_id_expired = entry_access_id_expired.get().strip()
    access_id_fw2 = entry_access_id_fw2.get().strip()
    access_id_expired3 = entry_access_id_expired3.get().strip()
    access_id_uid = entry_access_id_uid.get().strip()
    access_id_expired2 = entry_access_id_expired2.get().strip()
    sophos_case_ftp_username = entry_sophos_case_ftp_username.get().strip()
    sophos_case_ftp_password = entry_sophos_case_ftp_password.get().strip()
    frequency_of_the_issue = entry_frequency_of_the_issue.get().strip()
    repro_possible = combo_repro.get()
    repro_possible_details = entry_text_repro_possible.get("1.0", tk.END).strip()
    timestamp = entry_timestamp.get().strip()
    source_details = text_source_details.get("1.0", tk.END).strip()
    workaround_found = combo_workaround.get()
    workaround_desc = entry_workaround.get("1.0", tk.END).strip()
    poa = combo_poa.get()
    log_ftp = combo_log_ftp.get()
    log_ss = combo_log_ss.get()
    log_a = combo_log_a.get()
    
    # Checkbox values
    checkbox1_value = checkbox1_var.get()
    checkbox2_value = checkbox2_var.get()
    checkbox3_value = checkbox3_var.get()
    checkbox4_value = checkbox4_var.get()

    joined_path = os.path.join('C:/sophos-tmp', case_number)

    if not os.path.exists(joined_path):
        os.makedirs(joined_path)
    
    file_path = os.path.join(joined_path, f"SAP-{case_number}-{datetime.today().strftime('%Y-%m-%d')}.txt")
    with open(file_path, "w", encoding='utf-8') as external_file:
        external_file.write(f"***** SAP Case {case_number} from {datetime.today().strftime('%Y-%m-%d')} *****\n\n")
        
        if issue_description_cx:
            external_file.write(f"Short Issue Description = \n\n{issue_description_cx}\n\n")
        if issue_description_tse:
            external_file.write(f"Your Details = \n\n{issue_description_tse}\n\n")
        if error:
            external_file.write(f"Error message = \n\n{error}\n\n")
        
        external_file.write("***** DATA: ***** \n")
        if case_number:
            external_file.write(f"SAP Case = {case_number}\n")
        if customer_status:
            external_file.write(f"Status = {customer_status}\n")
        if customer_name:
            external_file.write(f"Customer / Partner Name = {customer_name}\n")
        if customer_phone:
            external_file.write(f"Customer / Partner Phone = {customer_phone}\n")
        if customer_language:
            external_file.write(f"Customer / Partner Language = {customer_language}\n")
        if new_configuration:
            external_file.write(f"New configuration = {new_configuration}\n")
        if product_cat:
           external_file.write(f"Product Category = {product_cat}\n")
        if product_typ:
            external_file.write(f"Product Type = {product_formatted_type}\n")
        if product_sn:
            external_file.write(f"Product serial number(s) = {product_sn}\n")
        if firmware_version:   
            external_file.write(f"Firmware / Software version = {firmware_version}\n")
        if ha:
            external_file.write(f"HA Cluster (yes/no) = {ha}\n")
        if ha_mode:
            external_file.write(f"High availability mode = {ha_mode}\n")
        if ha_count:
            external_file.write(f"How many devices in HA Cluster? = {ha_count}\n")
        if access_id_fw:
            external_file.write(f"Remote access ID firewall = {access_id_fw}\n")
        if access_id_expired:
            external_file.write(f"Remote access ID firewall expired = {access_id_expired}\n")
        if access_id_fw2:
            external_file.write(f"Remote access ID firewall2 = {access_id_fw2}\n")
        if access_id_expired3:
            external_file.write(f"Remote access ID firewall2 expired = {access_id_expired3}\n")
        if access_id_uid:
            external_file.write(f"Access UID central = {access_id_uid}\n")
        if access_id_expired2:
            external_file.write(f"Access UID central creation or expire Timestamp = {access_id_expired2}\n")
        if sophos_case_ftp_username or sophos_case_ftp_password:
            external_file.write(f"Sophos FTP server credentials = {sophos_case_ftp_username}:{sophos_case_ftp_password}\n\n")
        
        external_file.write("***** Issue Details: ***** \n")
        if frequency_of_the_issue:
            external_file.write(f"Issue Frequency = {frequency_of_the_issue}\n")
        if repro_possible:
            external_file.write(f"Repro possible = {repro_possible}\n")
        if repro_possible_details:
            external_file.write(f"Repro details explain = \n{repro_possible_details}\n\n")
        if timestamp:
            external_file.write(f"Timestamp = {timestamp}\n\n")
        if source_details:
            external_file.write(f"Source details = \n{source_details}\n\n")
        if workaround_found:
            external_file.write(f"Workaround found? Y/N = {workaround_found}\n")
        if workaround_desc:
            external_file.write(f"Workaround Details = \n{workaround_desc}\n\n")
        
        external_file.write("***** ASSESSMENT: ***** \n\n")
        external_file.write("***** TROUBLESHOOTING STEPS: ***** \n\n")
        external_file.write("***** ACTION PLAN: ***** \n")
        if poa:
            external_file.write(f"Plan of Action = {poa}\n\n")
        
        external_file.write("***** LOG ANALYSIS: ***** \n\n")
        external_file.write("***** LOGS: ***** \n\n")
        if log_ftp:
            external_file.write(f"> Logs on FTP = {log_ftp}\n")
        if log_ss:
            external_file.write(f"> Logs on SendSafely = {log_ss}\n")
        if log_a:
            external_file.write(f"> Logs on Attachment = {log_a}\n\n")
        
        # Write checkbox states
        external_file.write("***** CHECKBOXES: ***** \n")
        external_file.write(f"SOP used = {'YES' if checkbox1_value else 'NO'}\n")
        external_file.write(f"KB used = {'YES' if checkbox2_value else 'NO'}\n")
        external_file.write(f"Jira search used = {'YES' if checkbox3_value else 'NO'}\n")
        external_file.write(f"Issue resolved used = {'YES' if checkbox4_value else 'NO'}\n\n")
        
        external_file.write("***** FOLLOWED KB: ***** \n\n")
        external_file.write("***** COMMANDS USED FOR LOG COLLECTION: ***** \n\n\n\n")
        external_file.write("This SAP has been created with the SAP Note generator V3.0 GA \n\n")

    messagebox.showinfo("Success\a", f"SAP note created at {file_path}")

def import_file():
    file_path = askopenfilename(filetypes=[("Text files", "*.txt")])
    if not file_path:
        return

    with open(file_path, "r", encoding='utf-8') as file:
        content = file.readlines()

    fields = {
        "SAP Case": entry_case_number,
        "Short Issue Description": text_issue_description_cx,
        "Your Details": text_issue_description_tse,
        "Error message": text_entry_error,
        "Status": combo_costumer_status,
        "Customer / Partner Name": entry_customer_name,
        "Customer / Partner Phone": entry_customer_phone,
        "Customer / Partner Language": entry_customer_language,
        "New configuration": combo_new_configuration,
        "Product Category": combo_product_cat,
        "Product Type": entry_product_typ,
        "Product serial number(s)": entry_product_sn,
        "Firmware / Software version": entry_firmware_version,
        "HA Cluster (yes/no)": combo_ha,
        "High availability mode": combo_mode_ha,
        "How many devices in HA Cluster?": entry_ha_count,
        "Remote access ID firewall": entry_access_id_fw,
        "Remote access ID firewall expired": entry_access_id_expired,
        "Remote access ID firewall2": entry_access_id_fw2,
        "Remote access ID firewall2 expired": entry_access_id_expired3,
        "Access UID central": entry_access_id_uid,
        "Access UID central creation or expire Timestamp": entry_access_id_expired2,
        "Sophos FTP server credentials": (entry_sophos_case_ftp_username, entry_sophos_case_ftp_password),
        "Issue Frequency": entry_frequency_of_the_issue,
        "Repro possible": combo_repro,
        "Repro details explain": entry_text_repro_possible,
        "Timestamp": entry_timestamp,
        "Source details": text_source_details,
        "Workaround found? Y/N": combo_workaround,
        "Workaround Details": entry_workaround,
        "Plan of Action": combo_poa,
        "Logs on FTP": combo_log_ftp,
        "Logs on SendSafely": combo_log_ss,
        "Logs on Attachment": combo_log_a,
    }

    checkbox_fields = {
        "SOP used": checkbox1_var,
        "KB used": checkbox2_var,
        "Jira search used": checkbox3_var,
        "Issue resolved used": checkbox4_var,
    }

    key, value = None, None
    for line in content:
        if '=' in line:
            if key and value is not None:
                if isinstance(value, list):
                    value = '\n'.join(value)
                if key in fields:
                    widget = fields[key]
                    if isinstance(widget, tuple):  # For the FTP server credentials field
                        username, password = value.split(':')
                        widget[0].delete(0, tk.END)
                        widget[0].insert(0, username)
                        widget[1].delete(0, tk.END)
                        widget[1].insert(0, password)
                    elif isinstance(widget, tk.Entry):
                        widget.delete(0, tk.END)
                        widget.insert(0, value)
                    elif isinstance(widget, ScrolledText):
                        widget.delete("1.0", tk.END)
                        widget.insert(tk.END, value.replace('{}', "\n"))
                    elif isinstance(widget, ttk.Combobox):
                        widget.set(value)
                elif key in checkbox_fields:
                    checkbox_fields[key].set(value == 'YES')

            key, value = map(str.strip, line.split('=', 1))
            if key.startswith('> '):
                key = key[2:]
            if key in ["Short Issue Description", "Your Details", "Error message", "Repro details explain", "Source details", "Workaround Details"]:
                value = []
        elif key and isinstance(value, list):
            if "***** DATA: *****" in line and key == "Error message":
                value = '\n'.join(value)
                if key in fields:
                    widget = fields[key]
                    widget.delete("1.0", tk.END)
                    widget.insert(tk.END, value.replace('{}', "\n"))
                key, value = None, None
            elif "***** ASSESSMENT: *****" in line and key == "Workaround Details":
                value = '\n'.join(value)
                if key in fields:
                    widget = fields[key]
                    widget.delete("1.0", tk.END)
                    widget.insert(tk.END, value.replace('{}', "\n"))
                key, value = None, None
            else:
                value.append(line.strip())
        else:
            continue

    if key and value is not None:
        if isinstance(value, list):
            value = '\n'.join(value)
        if key in fields:
            widget = fields[key]
            if isinstance(widget, tuple):  # For the FTP server credentials field
                username, password = value.split(':')
                widget[0].delete(0, tk.END)
                widget[0].insert(0, username)
                widget[1].delete(0, tk.END)
                widget[1].insert(0, password)
            elif isinstance(widget, tk.Entry):
                widget.delete(0, tk.END)
                widget.insert(0, value)
            elif isinstance(widget, ScrolledText):
                widget.delete("1.0", tk.END)
                widget.insert(tk.END, value.replace('{}', "\n"))
            elif isinstance(widget, ttk.Combobox):
                widget.set(value)
        elif key in checkbox_fields:
            checkbox_fields[key].set(value == 'YES')

    messagebox.showinfo("Success", "Data imported successfully from the file.")

def export_to_word():
    file_path = asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if not file_path:
        return

    doc = Document()
    doc.add_heading(f"SAP Case {entry_case_number.get()} from {datetime.today().strftime('%Y-%m-%d')}", 0)

    def add_field(title, content):
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    add_field("Short Issue Description", text_issue_description_cx.get("1.0", tk.END).strip())
    add_field("Your Details", text_issue_description_tse.get("1.0", tk.END).strip())
    add_field("Error message", text_entry_error.get("1.0", tk.END).strip())
    
    doc.add_heading("DATA", level=1)
    add_field("Status", combo_costumer_status.get())
    add_field("Customer name", entry_customer_name.get().strip())
    add_field("Customer phone number", entry_customer_phone.get().strip())
    add_field("Customer language", entry_customer_language.get().strip())
    add_field("New configuration", combo_new_configuration.get())
    add_field("Product Category", combo_product_cat.get())
    add_field("Product Type", entry_product_typ.get().strip())
    add_field("Product serial number(s)", entry_product_sn.get().strip())
    add_field("Firmware / Software version", entry_firmware_version.get().strip())
    add_field("HA Cluster (yes/no)", combo_ha.get())
    add_field("High availability mode", combo_mode_ha.get())
    add_field("How many devices in HA Cluster?", entry_ha_count.get().strip())
    add_field("Remote access ID firewall", entry_access_id_fw.get().strip())
    add_field("Remote access ID firewall expired", entry_access_id_expired.get().strip())
    add_field("Remote access ID firewall2", entry_access_id_fw2.get().strip())
    add_field("Remote access ID firewall2 expired", entry_access_id_expired3.get().strip())
    add_field("Access UID central", entry_access_id_uid.get().strip())
    add_field("Access UID central creation or expire Timestamp", entry_access_id_expired2.get().strip())
    add_field("Sophos FTP server credentials", f"{entry_sophos_case_ftp_username.get().strip()}:{entry_sophos_case_ftp_password.get().strip()}")

    doc.add_heading("Issue Details", level=1)
    add_field("Issue Frequency", entry_frequency_of_the_issue.get().strip())
    add_field("Repro possible", combo_repro.get())
    add_field("Repro details explain", entry_text_repro_possible.get("1.0", tk.END).strip())
    add_field("Timestamp", entry_timestamp.get().strip())
    add_field("Source details", text_source_details.get("1.0", tk.END).strip())
    add_field("Workaround found? Y/N", combo_workaround.get())
    add_field("Workaround Details", entry_workaround.get("1.0", tk.END).strip())

    doc.add_heading("Troubleshooting", level=1)
    doc.add_paragraph(f"SOP used  {'YES' if checkbox1_var.get() else 'NO'}")
    doc.add_paragraph(f"KB used {'YES' if checkbox2_var.get() else 'NO'}")
    doc.add_paragraph(f"Jira search used {'YES' if checkbox3_var.get() else 'NO'}")
    doc.add_paragraph(f"Issue resolved used {'YES' if checkbox4_var.get() else 'NO'}")

    doc.save(file_path)
    messagebox.showinfo("Success", f"Word document saved at {file_path}")

def export_to_pdf():
    file_path = asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
    if not file_path:
        return

    c = Canvas(file_path, pagesize=letter)
    width, height = letter

    def draw_field(title, content, y):
        c.setFont("Helvetica-Bold", 12)
        c.drawString(30, y, title)
        c.setFont("Helvetica", 10)
        text_lines = content.split("\n")
        for line in text_lines:
            c.drawString(30, y-15, line)
            y -= 15
            if y < 50:  # Check if we need to add a new page
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - 50
        return y - 15

    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(30, y, f"SAP Case {entry_case_number.get()} from {datetime.today().strftime('%Y-%m-%d')}")
    y -= 40

    y = draw_field("Short Issue Description", text_issue_description_cx.get("1.0", tk.END).strip(), y)
    y = draw_field("Your Details", text_issue_description_tse.get("1.0", tk.END).strip(), y)
    y = draw_field("Error message", text_entry_error.get("1.0", tk.END).strip(), y)
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(30, y, "DATA")
    y -= 20

    y = draw_field("Status", combo_costumer_status.get(), y)
    y = draw_field("Customer name", entry_customer_name.get().strip(), y)
    y = draw_field("Customer phone number", entry_customer_phone.get().strip(), y)
    y = draw_field("Customer language", entry_customer_language.get().strip(), y)
    y = draw_field("New configuration", combo_new_configuration.get(), y)
    y = draw_field("Product", combo_product_cat.get(), y)
    y = draw_field("Product Type", entry_product_typ.get().strip(), y)
    y = draw_field("Product serial number(s)", entry_product_sn.get().strip(), y)
    y = draw_field("Firmware / Software version", entry_firmware_version.get().strip(), y)
    y = draw_field("High availability", combo_ha.get(), y)
    y = draw_field("High availability mode", combo_mode_ha.get(), y)
    y = draw_field("How many devices in HA Cluster?", entry_ha_count.get().strip(), y)
    y = draw_field("Remote access ID firewall", entry_access_id_fw.get().strip(), y)
    y = draw_field("Remote access ID firewall expired", entry_access_id_expired.get().strip(), y)
    y = draw_field("Remote access ID firewall2", entry_access_id_fw2.get().strip(), y)
    y = draw_field("Remote access ID firewall2 expired", entry_access_id_expired3.get().strip(), y)
    y = draw_field("Access UID central", entry_access_id_uid.get().strip(), y)
    y = draw_field("Access UID central creation or expire Timestamp", entry_access_id_expired2.get().strip(), y)
    y = draw_field("Sophos FTP server credentials", f"{entry_sophos_case_ftp_username.get().strip()}:{entry_sophos_case_ftp_password.get().strip()}", y)

    c.setFont("Helvetica-Bold", 12)
    c.drawString(30, y, "Issue Details")
    y -= 20

    y = draw_field("Issue Frequency", entry_frequency_of_the_issue.get().strip(), y)
    y = draw_field("Repro possible", combo_repro.get(), y)
    y = draw_field("Repro details explain", entry_text_repro_possible.get("1.0", tk.END).strip(), y)
    y = draw_field("Timestamp", entry_timestamp.get().strip(), y)
    y = draw_field("Source details", text_source_details.get("1.0", tk.END).strip(), y)
    y = draw_field("Workaround found? Y/N", combo_workaround.get(), y)
    y = draw_field("Workaround Details", entry_workaround.get("1.0", tk.END).strip(), y)

    c.setFont("Helvetica-Bold", 12)
    c.drawString(30, y, "Troubleshooting")
    y -= 20

    c.setFont("Helvetica", 10)
    c.drawString(30, y, f"SOP used = {'YES' if checkbox1_var.get() else 'NO'}")
    y -= 15
    c.drawString(30, y, f"KB used = {'YES' if checkbox2_var.get() else 'NO'}")
    y -= 15
    c.drawString(30, y, f"Jira search used = {'YES' if checkbox3_var.get() else 'NO'}")
    y -= 15
    c.drawString(30, y, f"Issue resolved used = {'YES' if checkbox4_var.get() else 'NO'}")

    c.showPage()
    c.save()
    messagebox.showinfo("Success", f"PDF document saved at {file_path}")

root = tk.Tk()
root.title("Sophos SAP Note File Generator by Michael Karek®")
root.geometry("650x750")

# Set the background color to #2196f3
bg_color = "#2196f3"
root.configure(bg=bg_color)

frame = tk.Frame(root, bg=bg_color)
frame.pack(fill="both", expand=True)

canvas = tk.Canvas(frame, bg=bg_color)
canvas.pack(side="left", fill="both", expand=True)

scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
scrollbar.pack(side="right", fill="y")

canvas.configure(yscrollcommand=scrollbar.set)
canvas.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
canvas.bind_all("<MouseWheel>", _on_mousewheel)

scrollable_frame = tk.Frame(canvas, bg=bg_color)
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

font = ("Arial", 12)

vcmd_int = (root.register(validate_int_input), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W')
vcmd_ver = (root.register(validate_version_input), '%d', '%i', '%P', '%s', '%S', '%v', '%V', '%W')

labels = [
    "Case Number*", "Short Issue Description", "Your Details", "Error Message", "Status",
    "Customer / Partner Name", "Customer / Partner Phone", "Customer / Partner Language", "New Configuration (yes/no)", 
    "Product Category", "Product Type\n (e.g. XG310, SG85W etc)", "Product Serial Number(s)", "Firmware / Software Version", 
    "HA Cluster (yes/no)", "HA Mode (AA/AP)", "HA Device Count", "Access ID Firewall", 
    "Access ID Firewall Expired Timestamp","Access ID Firewall2", 
    "Access ID Firewall2 Expired Timestamp", "Access UID Central", 
    "Access UID Central Expire Timestamp", "Sophos FTP Username", 
    "Sophos FTP Password", "Issue Frequency", "Reproducible (yes/no)", "Repro Details", "Repro Timestamp", 
    "Source Details", "Workaround (yes/no)", "Workaround Details", "POA", "Files on FTP", "Files on SendSafely", "Files on Attachment"
]

entries = []
combos = []

for i, label in enumerate(labels):
    tk.Label(scrollable_frame, text=label, font=font, bg=bg_color).grid(row=i, column=0, padx=10, pady=5, sticky="e")
    if label in ["Short Issue Description", "Your Details", "Error Message", "Source Details", "Repro Details", "Workaround Details"]:
        text_box = ScrolledText(scrollable_frame, height=5, width=30, font=font, bg="white")
        text_box.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        entries.append(text_box)
    elif label == "Case Number*":
        entry = tk.Entry(scrollable_frame, validate='key', validatecommand=vcmd_int, font=font, bg="white")
        entry.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        entries.append(entry)
    elif label in ["Status"]:
        combo = ttk.Combobox(scrollable_frame, values=[" ", "Customer", "Partner", "Partner Silver", "Partner Gold", "Partner Platinum", "Public Customer"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("Partner")
        combos.append(combo)
    elif label in ["New Configuration (yes/no)"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    elif label in ["Product Category"]:
        combo = ttk.Combobox(scrollable_frame, values=["CM/CFR", "Other", "Phish Threat", "Sophos Connect", "Sophos Email", "Sophos Switch", "UTM", "XG - 19.x", "XG - 20.x", "XG - 21.x", "ZTNA"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("XG - 20.x")
        combos.append(combo)
    elif label == "Firmware / Software Version":
        entry = tk.Entry(scrollable_frame, validate='key', validatecommand=vcmd_ver, font=font, bg="white")
        entry.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        entries.append(entry)
    elif label in ["HA Cluster (yes/no)"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    elif label in ["HA Mode (AA/AP)"]:
        combo = ttk.Combobox(scrollable_frame, values=["AA", "AP", ""], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("")
        combos.append(combo)
    elif label == "HA Device Count":
        entry = tk.Entry(scrollable_frame, validate='key', validatecommand=vcmd_int, font=font, bg="white")
        entry.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        entries.append(entry) 
    elif label in ["Reproducible (yes/no)"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    elif label in ["Workaround (yes/no)"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    elif label in ["POA"]:
        combo = ttk.Combobox(scrollable_frame, values=[" ", "Mentoring L2", "Mentoring GES", "GES Triage 1", "GES Triage 2", "GES Triage 3", "Triage CPG", "Move to Queue", "RMA", "Move to another TEAM"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set(" ")
        combos.append(combo)
    elif label in ["Files on FTP"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    elif label in ["Files on SendSafely"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    elif label in ["Files on Attachment"]:
        combo = ttk.Combobox(scrollable_frame, values=["YES", "NO"], font=font)
        combo.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        combo.set("NO")
        combos.append(combo)
    else:
        entry = tk.Entry(scrollable_frame, font=font, bg="white")
        entry.grid(row=i, column=1, padx=10, pady=5, sticky="w")
        entries.append(entry)

(
    entry_case_number, text_issue_description_cx, text_issue_description_tse, text_entry_error,
    entry_customer_name, entry_customer_phone, entry_customer_language, 
    entry_product_typ, entry_product_sn, entry_firmware_version,
    entry_ha_count, entry_access_id_fw, entry_access_id_expired,entry_access_id_fw2, 
    entry_access_id_expired3, entry_access_id_uid, entry_access_id_expired2, entry_sophos_case_ftp_username,
    entry_sophos_case_ftp_password, entry_frequency_of_the_issue, entry_text_repro_possible,
    entry_timestamp, text_source_details, entry_workaround
) = entries

(
combo_costumer_status, combo_new_configuration, combo_product_cat, combo_ha, combo_mode_ha, combo_repro, 
combo_workaround, combo_poa, combo_log_ftp, combo_log_ss, combo_log_a
)  = combos

# Create checkboxes
checkbox1_var = tk.BooleanVar()
checkbox2_var = tk.BooleanVar()
checkbox3_var = tk.BooleanVar()
checkbox4_var = tk.BooleanVar()

tk.Checkbutton(scrollable_frame, text="SOP USED", variable=checkbox1_var, font=font, bg=bg_color).grid(row=i+1, column=1, padx=10, pady=5, sticky="w")
tk.Checkbutton(scrollable_frame, text="KB Used", variable=checkbox2_var, font=font, bg=bg_color).grid(row=i+2, column=1, padx=10, pady=5, sticky="w")
tk.Checkbutton(scrollable_frame, text="Jira search", variable=checkbox3_var, font=font, bg=bg_color).grid(row=i+3, column=1, padx=10, pady=5, sticky="w")
tk.Checkbutton(scrollable_frame, text="Issue resolved", variable=checkbox4_var, font=font, bg=bg_color).grid(row=i+4, column=1, padx=10, pady=5, sticky="w")

tk.Button(scrollable_frame, text="SAP creation", command=submit, font=font).grid(row=len(labels)+4, columnspan=2, pady=10)
tk.Button(scrollable_frame, text="Import SAP", command=import_file, font=font).grid(row=len(labels)+5, columnspan=2, pady=10)
tk.Button(scrollable_frame, text="Export to Word", command=export_to_word, font=font).grid(row=len(labels)+5, columnspan=2, column=0, pady=10, padx=5, sticky="e")
tk.Button(scrollable_frame, text="Export to PDF", command=export_to_pdf, font=font).grid(row=len(labels)+5, columnspan=2, column=0, padx=10, pady=5, sticky="w")
tk.Button(scrollable_frame, text="Delete Entries", command=clearTextInput, font=font).grid(row=len(labels)+9, columnspan=2, column=0, padx=10, pady=5, sticky="w")
tk.Button(scrollable_frame, text="Quit", command=root.quit, font=font).grid(row=len(labels)+9, columnspan=2, pady=10, sticky="e")
root.mainloop()
